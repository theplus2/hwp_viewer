"""
HWP 파일에서 텍스트를 추출하는 모듈
hwp5 라이브러리를 사용하여 텍스트 추출
"""
import sys
import os
import tempfile
import html as html_lib
from typing import Tuple, List, Dict


def extract_text(file_path: str) -> str:
    """HWP/HWPX 파일에서 순수 텍스트 추출 (검색 인덱싱용) - 표 내용 포함"""
    
    # hwpx 파일은 별도 처리 (ZIP+XML 형식)
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.hwpx':
        return extract_text_from_hwpx(file_path)
    
    text = ""
    
    # 먼저 hwp5txt 시도 (가장 빠름)
    try:
        from hwp5.hwp5txt import main as hwp5txt_main
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as tmp:
            tmp_path = tmp.name
        
        saved_argv = sys.argv
        try:
            sys.argv = ['hwp5txt', '--output', tmp_path, file_path]
            hwp5txt_main()
            
            if os.path.exists(tmp_path) and os.path.getsize(tmp_path) > 0:
                with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read().strip()
        except SystemExit:
            pass
        except Exception:
            pass
        finally:
            sys.argv = saved_argv
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except:
                    pass
    except Exception:
        pass
    
    # hwp5txt로 추출한 텍스트가 너무 짧으면 xmlmodel로 표 내용도 추출
    # (표만 있는 문서의 경우 hwp5txt는 거의 빈 텍스트를 반환함)
    if len(text) < 50:
        try:
            text_with_tables = _extract_text_with_tables(file_path)
            if len(text_with_tables) > len(text):
                text = text_with_tables
        except Exception:
            pass
    
    return text


def extract_text_from_hwpx(file_path: str) -> str:
    """HWPX 파일에서 텍스트 추출 (ZIP+XML 형식)"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    text_parts = []
    
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            # Contents 폴더 내의 section*.xml 파일들에서 텍스트 추출
            for name in sorted(zf.namelist()):
                if name.startswith('Contents/section') and name.endswith('.xml'):
                    try:
                        xml_content = zf.read(name).decode('utf-8')
                        section_text = _extract_text_from_hwpx_section(xml_content)
                        if section_text:
                            text_parts.append(section_text)
                    except Exception:
                        pass
    except Exception:
        pass
    
    return '\n'.join(text_parts).strip()


def _extract_text_from_hwpx_section(xml_content: str) -> str:
    """HWPX section XML에서 텍스트 추출"""
    import xml.etree.ElementTree as ET
    
    text_parts = []
    
    try:
        # XML 네임스페이스 처리
        # HWPX는 다양한 네임스페이스를 사용하므로 태그에서 네임스페이스 제거
        root = ET.fromstring(xml_content)
        
        # 모든 텍스트 노드 순회
        for elem in root.iter():
            # 태그 이름에서 네임스페이스 제거
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
            
            # 텍스트 노드 (hp:t 또는 t 태그)
            if tag == 't' and elem.text:
                text_parts.append(elem.text)
            
            # 단락 끝에서 줄바꿈
            elif tag == 'p':
                text_parts.append('\n')
    except Exception:
        pass
    
    return ''.join(text_parts).strip()


def extract_html_from_hwpx(file_path: str) -> Tuple[str, List[Dict]]:
    """HWPX 파일에서 HTML 콘텐츠 추출"""
    import zipfile
    import xml.etree.ElementTree as ET
    
    html_parts = []
    
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            for name in sorted(zf.namelist()):
                if name.startswith('Contents/section') and name.endswith('.xml'):
                    try:
                        xml_content = zf.read(name).decode('utf-8')
                        section_html = _extract_html_from_hwpx_section(xml_content)
                        if section_html:
                            html_parts.append(section_html)
                    except Exception:
                        pass
    except Exception:
        pass
    
    if html_parts:
        return '\n'.join(html_parts), []
    
    return "<p>내용을 추출할 수 없습니다.</p>", []


def _extract_html_from_hwpx_section(xml_content: str) -> str:
    """HWPX section XML에서 HTML 추출 (구조적 파싱)"""
    import xml.etree.ElementTree as ET
    
    try:
        root = ET.fromstring(xml_content)
        return _parse_nodes(root)
    except Exception as e:
        print(f"HWPX Parsing Error: {e}")
        return ""


def _get_tag(elem) -> str:
    """XML 태그 이름에서 네임스페이스 제거"""
    return elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag


def _parse_nodes(elem) -> str:
    """XML 노드 순회하며 HTML 생성 (재귀)"""
    parts = []
    if elem is None:
        return ""
    
    for child in elem:
        tag = _get_tag(child)
        
        if tag == 'p':
            parts.append(_parse_paragraph(child))
        elif tag == 'tbl':
            parts.append(_parse_table(child))
        elif tag == 't':
            # 텍스트 노드
            text = child.text or ''
            parts.append(html_lib.escape(text))
        else:
            # run, sec, subList, tc 등 컨테이너 -> 재귀 호출
            parts.append(_parse_nodes(child))
            
    return ''.join(parts)


def _parse_paragraph(p_elem) -> str:
    """문단(p) 파싱"""
    content = _parse_nodes(p_elem)
    
    if not content.strip():
        return ""
        
    # 내용 중에 테이블이 있으면 p 태그로 감싸지 않는 것이 안전할 수 있음
    # (하지만 뷰어 호환성을 위해 스타일 조정 등으로 처리 가능, 여기서는 단순화)
    if '<table' in content:
        return content
        
    return f"<p>{content}</p>"


def _parse_table(tbl_elem) -> str:
    """표(tbl) 파싱"""
    rows_html = []
    
    # 자식 노드 중 tr 찾기
    for child in tbl_elem:
        if _get_tag(child) == 'tr':
            cells_html = []
            for tc in child:
                if _get_tag(tc) == 'tc':
                    # 셀(tc) 내용 재귀 파싱
                    # tc 안에 subList > p 구조가 일반적임
                    cell_content = _parse_nodes(tc)
                    cells_html.append(f'<td>{cell_content}</td>')
            
            if cells_html:
                rows_html.append('<tr>' + ''.join(cells_html) + '</tr>')
    
    if not rows_html:
        return ""
        
    return '<table border="1" style="border-collapse: collapse; width: 100%; margin: 10px 0;">' + ''.join(rows_html) + '</table>'


def _build_table_html_simple(rows: list) -> str:
    """(Deprecated) 간단한 테이블 HTML 생성 - 하위 호환성 유지용"""
    html = ['<table>']
    for row in rows:
        html.append('<tr>')
        for cell in row:
            escaped = html_lib.escape(cell) if cell else ''
            html.append(f'<td>{escaped}</td>')
        html.append('</tr>')
    html.append('</table>')
    return '\n'.join(html)


def _extract_text_with_tables(file_path: str) -> str:
    """xmlmodel을 사용하여 표 내용 포함 순수 텍스트 추출"""
    from hwp5.xmlmodel import Hwp5File
    from hwp5.treeop import STARTEVENT, ENDEVENT
    
    hwp = Hwp5File(file_path)
    text_parts = []
    
    try:
        for section in hwp.bodytext.sections:
            for event, item in section.events():
                model, attributes, context = item
                model_name = model.__name__
                
                # 텍스트 이벤트에서 텍스트 추출
                if model_name == 'Text' and event is STARTEVENT:
                    text = attributes.get('text', '')
                    if text:
                        text_parts.append(text)
                
                # 단락 끝에서 줄바꿈 추가
                elif model_name == 'Paragraph' and event is ENDEVENT:
                    text_parts.append('\n')
    except Exception:
        pass
    finally:
        hwp.close()
    
    return ''.join(text_parts).strip()



def extract_html(file_path: str) -> Tuple[str, List[Dict]]:
    """HWP/HWPX 파일에서 HTML 콘텐츠 추출 (표 내용 포함)"""
    
    # hwpx 파일은 별도 처리
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.hwpx':
        return extract_html_from_hwpx(file_path)
    
    # 먼저 xmlmodel로 표 포함 추출 시도
    try:
        html_content = _extract_with_xmlmodel(file_path)
        if html_content:
            return html_content, []
    except Exception:
        pass
    
    # 실패시 hwp5txt 폴백
    text = extract_text(file_path)
    if text:
        html_content = _text_to_html(text)
        return html_content, []
    
    return "<p>내용을 추출할 수 없습니다.</p>", []


def _extract_with_xmlmodel(file_path: str) -> str:
    """xmlmodel을 사용하여 표 내용까지 추출 (본문만, 중첩 표 지원)"""
    from hwp5.xmlmodel import Hwp5File
    from hwp5.treeop import STARTEVENT, ENDEVENT
    
    hwp = Hwp5File(file_path)
    html_parts = []
    
    # 표 관련 상태 (스택으로 중첩 표 지원)
    table_stack = []  # 각 요소: {'rows': [], 'current_row': []}
    in_cell = False
    current_cell_text = []
    
    # 단락 관련 상태 (텍스트를 단락 단위로 모음)
    in_paragraph = False
    paragraph_text = []
    
    try:
        # bodytext만 순회하여 머리글/바닥글 중복 방지
        for section in hwp.bodytext.sections:
            for event, item in section.events():
                model, attributes, context = item
                model_name = model.__name__
                
                # 표 시작
                if model_name == 'TableControl':
                    if event is STARTEVENT:
                        # 표 시작 전 단락 텍스트 저장
                        if paragraph_text and not table_stack:
                            text = ''.join(paragraph_text).strip()
                            if text:
                                escaped = html_lib.escape(text)
                                html_parts.append(f"<p>{escaped}</p>")
                            paragraph_text = []
                        # 새 표 시작 - 스택에 푸시
                        table_stack.append({'rows': [], 'current_row': []})
                    elif event is ENDEVENT:
                        # 표 종료 - 스택에서 팝하여 HTML 생성
                        if table_stack:
                            table_data = table_stack.pop()
                            if table_data['rows']:
                                table_html = _build_table_html(table_data['rows'])
                                if table_stack:
                                    # 중첩 표: 부모 표의 현재 셀에 추가
                                    current_cell_text.append(table_html)
                                else:
                                    # 최상위 표
                                    html_parts.append(table_html)
                
                # 표 행
                elif model_name == 'TableRow':
                    if event is STARTEVENT and table_stack:
                        table_stack[-1]['current_row'] = []
                    elif event is ENDEVENT and table_stack:
                        table_stack[-1]['rows'].append(table_stack[-1]['current_row'])
                
                # 표 셀
                elif model_name == 'TableCell':
                    if event is STARTEVENT:
                        in_cell = True
                        current_cell_text = []
                    elif event is ENDEVENT and table_stack:
                        cell_content = ''.join(current_cell_text)
                        table_stack[-1]['current_row'].append(cell_content)
                        in_cell = False
                        current_cell_text = []
                
                # 단락 시작/끝
                elif model_name == 'Paragraph':
                    if event is STARTEVENT:
                        in_paragraph = True
                    elif event is ENDEVENT:
                        in_paragraph = False
                        if in_cell and current_cell_text:
                            # 표 셀 내 단락 끝 - 줄바꿈 추가
                            current_cell_text.append('\n')
                        elif not table_stack and paragraph_text:
                            # 일반 단락 끝 - HTML 추가
                            text = ''.join(paragraph_text).strip()
                            if text:
                                escaped = html_lib.escape(text)
                                html_parts.append(f"<p>{escaped}</p>")
                            paragraph_text = []
                
                # 텍스트
                elif model_name == 'Text' and event is STARTEVENT:
                    text = attributes.get('text', '')
                    if in_cell:
                        current_cell_text.append(text)
                    elif not table_stack:
                        # 표 바깥의 일반 텍스트 - 단락에 누적
                        paragraph_text.append(text)
                
    except Exception:
        pass
    finally:
        hwp.close()
    
    # 마지막 단락 처리
    if paragraph_text:
        text = ''.join(paragraph_text).strip()
        if text:
            escaped = html_lib.escape(text)
            html_parts.append(f"<p>{escaped}</p>")
    
    return '\n'.join(html_parts)




def _build_table_html(rows: list) -> str:
    """표 행 데이터를 HTML 테이블로 변환"""
    html = ['<table>']
    for row in rows:
        html.append('<tr>')
        for cell in row:
            if cell:
                # 중첩 표인 경우 이스케이프하지 않음
                if cell.strip().startswith('<table>'):
                    html.append(f'<td>{cell}</td>')
                else:
                    escaped = html_lib.escape(cell)
                    # 줄바꿈을 <br> 태그로 변환
                    escaped = escaped.replace('\n', '<br>')
                    html.append(f'<td>{escaped}</td>')
            else:
                html.append('<td></td>')
        html.append('</tr>')
    html.append('</table>')
    return '\n'.join(html)




def _text_to_html(text: str) -> str:
    """텍스트를 HTML로 변환"""
    lines = text.split('\n')
    html_parts = []
    
    for line in lines:
        line = line.strip()
        if line:
            escaped = html_lib.escape(line)
            html_parts.append(f"<p>{escaped}</p>")
    
    return '\n'.join(html_parts)



# DOCX 지원
def extract_text_from_docx(file_path: str) -> str:
    """DOCX 파일에서 텍스트 추출"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception:
        return ""


def extract_html_from_docx(file_path: str) -> Tuple[str, List[Dict]]:
    """DOCX 파일에서 HTML 콘텐츠 및 이미지 추출"""
    import zipfile
    import base64
    
    html_parts = []
    images = []
    MAX_IMAGE_SIZE = 500 * 1024  # 500KB 제한
    MAX_IMAGES = 10  # 최대 10개 이미지
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        # 이미지 추출 (DOCX는 ZIP 파일)
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                for name in zf.namelist():
                    if len(images) >= MAX_IMAGES:
                        break
                    if name.startswith('word/media/'):
                        try:
                            img_data = zf.read(name)
                            # 크기 제한 확인
                            if len(img_data) > MAX_IMAGE_SIZE:
                                continue
                            
                            img_name = os.path.basename(name)
                            ext = os.path.splitext(img_name)[1].lower()
                            
                            # MIME 타입 결정
                            mime_types = {
                                '.png': 'image/png',
                                '.jpg': 'image/jpeg',
                                '.jpeg': 'image/jpeg',
                                '.gif': 'image/gif',
                                '.bmp': 'image/bmp'
                            }
                            mime = mime_types.get(ext, 'image/png')
                            
                            images.append({
                                'name': img_name,
                                'data': base64.b64encode(img_data).decode('utf-8'),
                                'mime': mime
                            })
                        except Exception:
                            continue
        except Exception:
            pass  # 이미지 추출 실패 시 텍스트만 표시
        
        # 문단 추출
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                escaped = html_lib.escape(text)
                html_parts.append(f"<p>{escaped}</p>")
        
        # 이미지 삽입 (크기 제한된 이미지만)
        for img in images:
            html_parts.append(f'<img src="data:{img["mime"]};base64,{img["data"]}" style="max-width: 100%; max-height: 400px; margin: 10px 0;" alt="{img["name"]}">')
        
        # 표 추출
        for table in doc.tables:
            html_parts.append(_table_to_html(table))
        
        if html_parts:
            return '\n'.join(html_parts), []
            
    except Exception as e:
        return f"<p>DOCX 파일 읽기 오류: {str(e)}</p>", []
    
    return "<p>내용을 추출할 수 없습니다.</p>", []


def _table_to_html(table) -> str:
    """DOCX 표를 HTML 테이블로 변환"""
    html = ['<table>']
    for row in table.rows:
        html.append('<tr>')
        for cell in row.cells:
            text = html_lib.escape(cell.text.strip())
            html.append(f'<td>{text}</td>')
        html.append('</tr>')
    html.append('</table>')
    return '\n'.join(html)


# TXT 지원
def extract_text_from_txt(file_path: str) -> str:
    """TXT 파일에서 텍스트 추출"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception:
        return ""


def extract_html_from_txt(file_path: str) -> Tuple[str, List[Dict]]:
    """TXT 파일에서 HTML 콘텐츠 추출"""
    text = extract_text_from_txt(file_path)
    if text:
        return _text_to_html(text), []
    return "<p>내용을 추출할 수 없습니다.</p>", []
